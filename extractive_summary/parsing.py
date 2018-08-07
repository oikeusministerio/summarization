
import numpy as np
from docx import Document
from nltk import sent_tokenize
import pandas as pd
import re
from tools.tools import sentence_tokenize
import textract
from tempfile import NamedTemporaryFile
import os

def count_sentences_left(sections):
    sections_df = pd.Series(np.array(sections))
    count_sentences = lambda x: len([s for s in sent_tokenize(x, language="finnish") if len(s) > 2])
    sentence_lengths = sections_df.apply(count_sentences)
    sentences_left = sentence_lengths.iloc[::-1].cumsum().iloc[::-1]
    return sentence_lengths, sentences_left

def split_too_long_sections(parsed_document, titles, sections_max_length, section_min_sentence):
    # let's check if some section is too long
    for title_i,title in enumerate(titles):
        section = parsed_document[title]
        sents = sentence_tokenize(section)
        n = len(sents)
        if n > sections_max_length:
            optimal = (section_min_sentence + sections_max_length) / 2
            times = int(np.ceil(n / optimal))
            del parsed_document[title]
            del titles[title_i]
            last = 0
            for i in range(1, times + 1):
                new_title = title + "_" + str(i)
                titles.insert(title_i + i - 1, new_title)
                parsed_document[new_title] = " ".join(sents[last:np.min([int(optimal*i), n])])
                last = int(optimal*i)
    return parsed_document, titles

def replace_words_in_txt(parsed_document, titles, word_list, substitutes):
    """
    Replaces all occurances of words in text.
    :param word_list: list of all words to be replaced
    :param substitute: what to put instead of word in list
    :return:
    """
    replaced_document = {}
    for title in titles:
        replaced_document[title] = {}
        replaced_document[title]['summary'] = parsed_document[title]
        for i, w in enumerate(word_list):
            replaced_document[title]['summary'] = replaced_document[title]['summary'].replace(w, substitutes[i])
    return replaced_document

class DocumentParser:

    txt_section_split_by = "\n[ ]*\n"  # txt parsing uses this regex to find double newlines with space between
    txt_section_delimiter = "\n\n"  # for testing, for example this can be used as delimiter

    def __init__(self, file):
        self.file = file
        self.body_text_styles = ['Body Text','Normal']

    def parse_docx_file(self):
        """
        The following is supposed when parsing :
        1) if there are titles in the text, the text will start with a title and will never end with one
        2) if there are titles in the text, text before first title or heading is skipped
        :return:
        """
        self.document = Document(self.file)
        parsed_document = {}
        titles = []
        sections_starting = []
        for i, paragraph in enumerate(self.document.paragraphs):
            style = paragraph.style.name
            if style not in self.body_text_styles:
                titles.append(paragraph.text.strip())
                sections_starting.append(i)
        paragraphs = np.array([p.text for p in self.document.paragraphs])
        sections_starting.append(len(paragraphs))
        for i,title in enumerate(titles):
            pars = paragraphs[sections_starting[i]+1:sections_starting[i+1]]
            parsed_document[title] = " ".join(list(pars))

        if len(titles) == 0:
            parsed_document['Sisalto'] =  " ".join(list(paragraphs))
            titles.append('Sisalto')

        return parsed_document, titles

    def read_docx_document(self):
        """
        :return: full .docx file
        """
        return ' '.join([paragraph.text for paragraph  in Document(self.file).paragraphs])


    def parse_txt(self, text, section_min_sentence=50, sections_max_length=175):
        """
        File is splitted by two newlines, that might have spaces between them.

        :param length_by: either by sentence count or character count
        :param section_min_length:
        :param section_min_sentence:
        :return: parsed txt file
        """
        titles = self.get_titles(text)
        parsed_document = self.split_text_by_titles(text, titles)
        parsed_document['titles'] = titles
        parsed_document, titles = split_too_long_sections(parsed_document, titles, sections_max_length,
                                                          section_min_sentence)

        return parsed_document, titles

    def parse_txt_file(self, section_min_sentence=50, sections_max_length=175):
        """
        File is splitted by two newlines, that might have spaces between them.

        :param length_by: either by sentence count or character count
        :param section_min_length:
        :param section_min_sentence:
        :return: parsed txt file
        """
        self.text = self.file.read().decode('utf8')
        return self.parse_txt(self.text, section_min_sentence=section_min_sentence,sections_max_length=sections_max_length)

    def paragraph_count(self):
        return len(self.document.paragraphs)

    def read_txt_document(self):
        return self.file.read().decode('utf8')

    def read_pdf_document(self):
        with NamedTemporaryFile(suffix='.pdf') as tmp_file:
            filename = tmp_file.name

            self.file.save(filename)
            raw_text = textract.process(filename, layout=True)
            text_without_pagebreaks = raw_text.replace(b'\x0c', b' ')

            text = text_without_pagebreaks.decode('utf-8')
            callback = lambda pat: pat.group(0)[0] + ' ' + pat.group(0)[2]
            text = re.sub('[\w,\.-]\\n\w', callback, text)

            return text

    def parse_pdf_file(self, section_min_sentence=50, sections_max_length=175):
        text = self.read_pdf_document()
        titles = self.get_titles(text)

        parsed_document = self.split_text_by_titles(text, titles)

        parsed_document['titles'] = titles
        parsed_document, titles = split_too_long_sections(parsed_document, titles, sections_max_length,
                                                          section_min_sentence)
        return parsed_document, titles

    def get_titles(self, text):
        lines = [line.strip() for line in re.split('\n', text)]
        return [line for i, line in enumerate(lines)
                if (i >= 1 and lines[i - 1] == '' and len(line) > 0 and not line.endswith('.'))]

    def split_text_by_titles(self,text,titles):
        title_starts = np.array([text.find(title) for title in titles])
        title_lens = np.array([len(t) for t in titles])
        section_starts = title_starts + title_lens
        section_ends = np.append(title_starts[1:], np.array([len(text)]))

        parsed_document = {}
        for i in range(len(titles)):
            parsed_document[titles[i]] = text[section_starts[i]:section_ends[i]]
        return parsed_document