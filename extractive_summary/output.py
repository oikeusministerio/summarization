
from docx import Document
import tempfile

class SummaryWriter:

    def __init__(self, summaries):
        self.summaries = summaries

    def write_docx(self, destination_path):
        """
       Writes summaries to given file in docx format.
        :return:
        """
        document = Document()
        for filename in self.summaries['filenames']:
            document.add_heading(filename, level=1)
            sections = self.summaries[filename]
            for title in sections['titles']:
                document.add_heading(title, level=2)
                document.add_paragraph(sections[title]['summary'])


        document.save(destination_path)

    def write_txt(self, destination_path):
        """
       Writes summaries to given file in docx format.
        :return:
        """
        text = ''
        for filename in self.summaries['filenames']:
            text += filename + '\n\n\n'
            sections = self.summaries[filename]
            for title in sections['titles']:
                text += title + '\n\n'
                text += sections[title]['summary'] + '\n'


        with open(destination_path, 'w') as f:
            f.write(text)
