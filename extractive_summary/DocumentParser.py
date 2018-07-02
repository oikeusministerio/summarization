
import numpy as np
from docx import Document
from nltk import sent_tokenize
import pandas as pd
import re


def count_sentences_left(sections):
    sections_df = pd.Series(np.array(sections))
    count_sentences = lambda x: len([s for s in sent_tokenize(x, language="finnish") if len(s) > 2])
    sentence_lengths = sections_df.apply(count_sentences)
    sentences_left = sentence_lengths.iloc[::-1].cumsum().iloc[::-1]
    return sentence_lengths, sentences_left

def split_too_long_sections(parsed_document, titles, sections_max_length, section_min_sentence):
    # let's check if some section is too long
    for title_i,title in enumerate(titles):
        section = parsed_document[title]
        sents = sent_tokenize(section, language="finnish")
        n = len(sents)
        if n > sections_max_length:
            optimal = (section_min_sentence + sections_max_length) / 2
            times = int(np.ceil(n / optimal))
            del parsed_document[title]
            del titles[title_i]
            last = 0
            for i in range(1, times + 1):
                new_title = title + "_" + str(i)
                titles.insert(title_i + i - 1, new_title)
                parsed_document[new_title] = " ".join(sents[last:np.min([int(optimal*i), n])])
                last = int(optimal*i)
    return parsed_document, titles


class DocumentParser:

    txt_section_split_by = "\n[ ]*\n"  # txt parsing uses this regex to find double newlines with space between
    txt_section_delimiter = "\n\n"  # for testing, for example this can be used as delimiter

    def __init__(self, file):
        self.file = file
        self.body_text_styles = ['Body Text','Normal']

    def parse_docx(self):
        """
        The following is supposed when parsing :
        1) if there are titles in the text, the text will start with a title and will never end with one
        2) if there are titles in the text, text before first title or heading is skipped
        :return:
        """
        self.document = Document(self.file)
        parsed_document = {}
        titles = []
        sections_starting = []
        for i, paragraph in enumerate(self.document.paragraphs):
            style = paragraph.style.name
            if style not in self.body_text_styles:
                titles.append(paragraph.text.strip())
                sections_starting.append(i)
        paragraphs = np.array([p.text for p in self.document.paragraphs])
        sections_starting.append(len(paragraphs))
        for i,title in enumerate(titles):
            pars = paragraphs[sections_starting[i]+1:sections_starting[i+1]]
            parsed_document[title] = list(pars)

        if len(titles) == 0:
            parsed_document['Sisalto'] =  list(paragraphs)
            titles.append('Sisalto')

        return parsed_document, titles



    def parse_txt(self, section_min_sentence=50, sections_max_length=175):
        """
        File is splitted by two newlines, that might have spaces between them.

        :param length_by: either by sentence count or character count
        :param section_min_length:
        :param section_min_sentence:
        :return: parsed txt file
        """
        self.text = self.file.read().decode('utf8')
        total_sents = len(sent_tokenize(self.text, language="finnish"))
        assert total_sents > section_min_sentence, "Cannot parse text to sections of " + \
                                                   str(section_min_sentence) + ", there are only " + str(total_sents) + " in total."
        rx_seq = re.compile(DocumentParser.txt_section_split_by)
        sections = rx_seq.split(str(self.text))
        sections = [s for s in sections if len(s) > 2]
        sections_N = len(sections)
        parsed_document = {}
        headings_candidates = []
        titles = []

        is_heading = lambda x: len(sent_tokenize(x, language="finnish")) == 1
        for i, section in enumerate(sections[:-1]): # last one cannot be heading
            if is_heading(section):
                headings_candidates.append((i,section.strip()))

        sents_lengths, sents_left = count_sentences_left(sections)

        last_section_used = -1
        for heading_i, item in enumerate(headings_candidates):
            i,_ = item
            if i < last_section_used:
                continue
            j = 1
            section = sections[i + j]
            section_sents = sents_lengths[i + j]
            if sents_left[i + j + 1] < section_min_sentence: # this means there are not enough characters/sentences to make full section.
                while i + j + 2 < sections_N:
                    j += 1
                    section += sections[i + j]
                    section_sents += sents_lengths[i + j]
            else:
                while section_sents < section_min_sentence and i + j + 2 < sections_N:
                    j += 1
                    section += sections[i + j]
                    section_sents += sents_lengths[i + j]


            title = sections[i].strip()
            if title in parsed_document:
                title += '_'
            parsed_document[title] = section
            titles.append(title)
            last_section_used = i + j

        while last_section_used < sections_N - 1:
            last_section_used += 1
            parsed_document[titles[-1]] += sections[last_section_used]

        parsed_document, titles = split_too_long_sections(parsed_document,titles,sections_max_length,section_min_sentence)

        return parsed_document, titles

    def paragraph_count(self):
        return len(self.document.paragraphs)
